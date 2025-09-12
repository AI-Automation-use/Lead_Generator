import logging
import azure.functions as func
import requests
import pandas as pd
from datetime import datetime, timedelta
import os
from dotenv import load_dotenv
from playwright.sync_api import sync_playwright
from bs4 import BeautifulSoup
import sys
import asyncio
from openai import AzureOpenAI
from docx import Document 
from newspaper import Article
from urllib.parse import quote
import msal
import base64
import re
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient
from io import BytesIO
import io
from azure.identity import DefaultAzureCredential

load_dotenv()

# Constants for authentication
CLIENT_ID = os.getenv("CLIENT_ID")
TENANT_ID = os.getenv("TENANT_ID")
TOKEN_FILE = "token.json"
SCOPES = ["Mail.Send"]
GRAPH_API_ENDPOINT = "https://graph.microsoft.com/v1.0"

# Accessing the values from the .env file
API_KEY = os.getenv("API_KEY")
CX = os.getenv("CX")
AZURE_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")
GNEWS_API_KEY = os.getenv("GNEWS_API_KEY")

# Azure Blob Storage Configuration
AZURE_STORAGE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
LEAD_EXCEL_CONTAINER_NAME="potentiallist"
LEAD_EXCEL_BLOB_NAME = "leads_tracking.xlsx" 

# Set up Azure OpenAI API
client = AzureOpenAI(
    api_key=AZURE_API_KEY,
    api_version="2024-12-01-preview",
    azure_endpoint=AZURE_ENDPOINT
)

# Ensuring Windows event loop policy for Playwright
# NOTE: This part is for Windows local development and can be removed for Linux-based
# Azure Function deployments, but it won't cause an error if left in.
if sys.platform == "win32":
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

# GNews API Configuration for fetching news
BASE_URL = "https://gnews.io/api/v4/search"


def get_blob_service_client():
    return BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)

def download_excel_from_blob(blob_service_client, container_name, blob_name):
    try:
        container_client = blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(blob_name)
        download_stream = blob_client.download_blob()
        return io.BytesIO(download_stream.readall())
    except Exception as e:
        print(f"Error downloading blob {blob_name}: {e}")
        return None

def get_identified_leads_df():
    blob_service_client = get_blob_service_client()
    excel_data = download_excel_from_blob(blob_service_client, LEAD_EXCEL_CONTAINER_NAME, LEAD_EXCEL_BLOB_NAME)
    if excel_data:
        try:
            df = pd.read_excel(excel_data)
            return df
        except Exception as e:
            print(f"Error reading Excel from blob: {e}")
            return pd.DataFrame(columns=["Company Name", "Lead Identification Areas", "Timestamp"])
    else:
        return pd.DataFrame(columns=["Company Name", "Lead Identification Areas", "Timestamp"])

def normalize_areas_string(areas_str):
    if not isinstance(areas_str, str):
        return ""
    parts = areas_str.replace(';', ',').split(',')
    cleaned_parts = sorted(list(set(area.strip() for area in parts if area.strip())))
    return ", ".join(cleaned_parts)

def add_lead_to_excel(company_name, lead_areas):
    blob_service_client = get_blob_service_client()
    df = get_identified_leads_df()
    normalized_incoming_areas_str = normalize_areas_string(lead_areas)
    incoming_areas_set = set(normalized_incoming_areas_str.split(', ') if normalized_incoming_areas_str else set())
    existing_company_row = df[df["Company Name"] == company_name]
    email_should_be_sent = False
    if not existing_company_row.empty:
        existing_areas_str = existing_company_row["Lead Identification Areas"].iloc[0]
        existing_areas_set = set(existing_areas_str.split(', ') if existing_areas_str else set())
        truly_new_areas = incoming_areas_set - existing_areas_set
        if truly_new_areas:
            updated_areas_set = existing_areas_set.union(incoming_areas_set)
            updated_areas_list = sorted(list(updated_areas_set))
            updated_areas_str = ", ".join(updated_areas_list)
            df.loc[df["Company Name"] == company_name, "Lead Identification Areas"] = updated_areas_str
            df.loc[df["Company Name"] == company_name, "Timestamp"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            output = io.BytesIO()
            df.to_excel(output, index=False)
            output.seek(0)
            upload_excel_to_blob(blob_service_client, LEAD_EXCEL_CONTAINER_NAME, LEAD_EXCEL_BLOB_NAME, output.getvalue())
            print(f"Company '{company_name}' updated with new lead areas: {', '.join(sorted(list(truly_new_areas)))}.")
            email_should_be_sent = True
        else:
            print(f"Company '{company_name}' already exists with these lead areas. No update or email needed.")
            email_should_be_sent = False
    else:
        new_entry = pd.DataFrame([{
            "Company Name": company_name,
            "Lead Identification Areas": normalized_incoming_areas_str,
            "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }])
        df = pd.concat([df, new_entry], ignore_index=True)
        output = io.BytesIO()
        df.to_excel(output, index=False)
        output.seek(0)
        upload_excel_to_blob(blob_service_client, LEAD_EXCEL_CONTAINER_NAME, LEAD_EXCEL_BLOB_NAME, output.getvalue())
        print(f"New company '{company_name}' added as a lead with areas: {normalized_incoming_areas_str}.")
        email_should_be_sent = True
    return email_should_be_sent

def fetch_full_article_text_with_playwright(page, url):
    try:
        page.goto(url, wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(3000)
        content = page.locator("body").inner_text()
        clean_text = content.strip().replace('\n', ' ').replace('\r', ' ')
        return clean_text[:10000] if clean_text else "⚠️ Full article not available."
    except Exception as e:
        print(f"Error fetching full article from {url}: {e}")
        return "⚠️ Full article not available."

def scrape_google_news(company_name, pages=1):
    query = quote(company_name)
    results = []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(user_agent=(
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/113.0.0.0 Safari/537.36"
        ))
        page = context.new_page()
        for i in range(pages):
            start = i * 10
            url = f"https://www.google.com/search?q={query}&tbm=nws&start={start}"
            page.goto(url, wait_until="load", timeout=60000)
            soup = BeautifulSoup(page.content(), "html.parser")
            for result in soup.select("div.SoaBEf"):
                try:
                    a_tag = result.find("a", href=True)
                    link = a_tag["href"] if a_tag else ""
                    title_el = a_tag.select_one("div.n0jPhd.ynAwRc.MBeuO.nDgy9d")
                    description_el = a_tag.select_one("div.GI74Re.nDgy9d")
                    publisher_el = a_tag.select_one("span.xQ82C.e8fRJf")
                    date_el = result.select_one("span[class]:not([class*='xQ82C'])")
                    title = title_el.get_text(strip=True) if title_el else ""
                    short_description = description_el.get_text(strip=True) if description_el else ""
                    publisher = publisher_el.get_text(strip=True) if publisher_el else ""
                    published_on = date_el.get_text(strip=True) if date_el else ""
                    full_article = fetch_full_article_text_with_playwright(page, link)
                    if title and link and not any(r["url"] == link for r in results):
                        results.append({
                            "title": title,
                            "publisher": publisher,
                            "published_on": published_on,
                            "description": full_article or short_description,
                            "url": link
                        })
                except Exception:
                    continue
        browser.close()
    return results


def upload_excel_to_blob(blob_service_client, container_name, blob_name, data_stream):
    try:
        container_client = blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(blob_name)
        blob_client.upload_blob(data_stream, overwrite=True)
        print(f"Successfully uploaded {blob_name} to blob storage.")
    except Exception as e:
        print(f"Error uploading blob {blob_name}: {e}")


# Constants for Blob Storage
TOKEN_CONTAINER_NAME = "potentiallist"
TOKEN_BLOB_NAME = "token.json"
SCOPES = ["https://graph.microsoft.com/.default"] # Use the .default scope for client credential flow

def get_access_token():
    logging.info("🔐 Acquiring access token...")
    token_cache = msal.SerializableTokenCache()
    blob_service_client = BlobServiceClient.from_connection_string(os.getenv("AZURE_STORAGE_CONNECTION_STRING"))
    
    # 1. Download the token cache from blob storage
    try:
        container_client = blob_service_client.get_container_client(TOKEN_CONTAINER_NAME)
        blob_client = container_client.get_blob_client(TOKEN_BLOB_NAME)
        download_stream = blob_client.download_blob()
        token_cache.deserialize(download_stream.readall().decode('utf-8'))
        logging.info("✅ Loaded token from blob storage cache.")
    except Exception as e:
        logging.warning(f"⚠️ Failed to load token cache from blob. Initial authentication may be required: {e}")
        # Note: This is where a local interactive session is needed to create the initial token.json.
        # This code block will fail in a deployed Function App.
        # It's here for local testing only. After initial setup, the 'accounts' check below handles it.
    
    app = msal.PublicClientApplication(
        client_id=os.getenv("CLIENT_ID"),
        authority=f"https://login.microsoftonline.com/{os.getenv('TENANT_ID')}",
        token_cache=token_cache
    )
    
    accounts = app.get_accounts()
    
    # Use the refresh token from the cache (downloaded from blob)
    result = app.acquire_token_silent(SCOPES, account=accounts[0]) if accounts else None
    
    # If a silent token acquisition fails, it means the refresh token is expired or not present.
    # In a deployed function, this should NOT happen if the refresh token is valid.
    if not result:
        raise Exception("❌ Token acquisition failed silently. Refresh token may be invalid. Manual re-authentication is required.")
        
    if token_cache.has_state_changed:
        try:
            container_client = blob_service_client.get_container_client(TOKEN_CONTAINER_NAME)
            blob_client = container_client.get_blob_client(TOKEN_BLOB_NAME)
            blob_client.upload_blob(token_cache.serialize(), overwrite=True)
            logging.info("💾 Token cache updated and saved to blob storage.")
        except Exception as e:
            logging.error(f"❌ Failed to upload token cache to blob: {e}")
            
    if "access_token" not in result:
        raise Exception(f"❌ Token acquisition failed: {result.get('error_description')}")
        
    return result["access_token"]

def get_company_website(company_name, api_key, cx):
    search_url = f"https://www.googleapis.com/customsearch/v1?q={company_name}+company+site&key={api_key}&cx={cx}"
    response = requests.get(search_url)
    if response.status_code == 200:
        results = response.json()
        if 'items' in results:
            return results['items'][0]['link']
    return None

def scrape_website(website):
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            ignore_https_errors=True
        )
        page = context.new_page()
        page.set_extra_http_headers({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9",
            "Connection": "keep-alive",
        })
        page.goto(website, wait_until="load", timeout=60000)
        page_content = page.content()
        soup = BeautifulSoup(page_content, 'html.parser')
        title = soup.title.string if soup.title else "No title found"
        paragraphs = soup.find_all('p')
        paragraphs_content = '\n'.join([para.get_text() for para in paragraphs])
        headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        headings_content = '\n'.join([heading.get_text() for heading in headings])
        full_content = f"**Title:** {title}\n\n"
        full_content += f"**Headings:**\n{headings_content}\n\n"
        full_content += f"**Paragraphs:**\n{paragraphs_content}\n\n"
        full_content += f"\n\n**Source:** {website}"
        browser.close()
        return full_content, website

def check_potential_lead(website_content, linkedin_content, news_content):
    combined_content = f"""
    You are a market-intelligence assistant.

    Analyze the following source materials (Website, LinkedIn, News) and produce a text report with these three sections.
    
    Decision rules (strict):
    - For each area, answer “Yes” ONLY if there is CLEAR, COMPANY-SPECIFIC evidence of current interest or activity in that area within the last 3 days found in the Sources below. Examples: official announcements, press releases, case studies, product/initiative pages tied to THIS company, exec statements, RFPs/tenders, partnerships, hiring/posts explicitly for that area, migration/adoption milestones.
    - Generic capability pages, vague marketing language, industry articles not tied to THIS company, or historical items older than 3 days → “No”.
    - Use ONLY the provided Sources block. Do NOT invent links or content.
    - For every “Yes”, include at least one exact excerpt in the Evidence section PLUS the source URL and (if available) the date. If you cannot provide an excerpt + URL from the Sources block, answer “No”.
    
    1. **Interest**
       For each area below, state “Yes” or “No”. If “Yes”, add a one-sentence why with an inline source tag like (Website), (LinkedIn), or (News).
    
       New Business Development Areas:
         - SFR150
         - Zones
         - DYN365
         - AI
         - AWS
    
       Large Deal Areas:
         - Cost Takeout
         - Cloud Migration
         - Data Migration
         - Platform Migration
         - SaaS
         - GCC
         - Partner with IT
    
    2. **Contacts**
       List each contact found, with type (email/phone/name+title) and the source (Website, LinkedIn, or News). Use only items present in Sources.
    
    3. **Evidence**
       Under sub-headings for Website, LinkedIn, and News, give the exact excerpt(s) that support each “Yes” or any contacts found. For each excerpt include the source URL and (if present) the date. If no supporting excerpts exist for a given area, do not fabricate; that area must be “No”.
    
    **Sources**
    Website Content:
    {website_content}
    
    LinkedIn Profile Content:
    {linkedin_content}
    
    Recent News Articles:
    {news_content}
    
    """
    response = client.chat.completions.create(
        model=AZURE_DEPLOYMENT,
        messages=[{"role": "user", "content": combined_content}],
    )
    lead_analysis = response.choices[0].message.content
    potential_lead_check = classify_lead(lead_analysis)
    return lead_analysis, potential_lead_check


def classify_lead(lead_analysis):
    prompt = f"""
    You are a proactive lead generation expert.

    Given the following lead analysis, determine if this company shows any sign—direct or indirect—of being a potential lead. Even a slight indication of interest, relevance, or alignment should result in "Yes".

    **Lead Analysis:**
    {lead_analysis}

    Answer strictly with "Yes" or "No".
    """
    response = client.chat.completions.create(
        model=AZURE_DEPLOYMENT,
        messages=[{"role": "user", "content": prompt}],
    )
    result = response.choices[0].message.content.strip()
    return result

def extract_lead_details(lead_analysis, company):
    prompt = f"""
    Given the following lead analysis for the company "{company}", extract and return the following in plain text format:
    
    1. Customer Name
    2. Lead Identification Area(s) (e.g., SFR150, Zones, DYN365, AI, AWS, Cost Takeout, Cloud Migration, Data Migration, Platform Migration, SaaS, GCC, Partner with IT)

    Format the output like this:
    
    **Customer Name**: <value>
    **Lead Identification Area**: <value>
    
    If a field is not found, just leave it like.

    Lead Analysis:
    {lead_analysis}
    """
    response = client.chat.completions.create(
        model=AZURE_DEPLOYMENT,
        messages=[{"role": "user", "content": prompt}],
    )
    extracted = response.choices[0].message.content.strip()
    return extracted

def create_lead_docx(lead_analysis: str, company: str):
    doc = Document()
    doc.add_heading(f"Lead Analysis for {company}", 0)
    doc.add_paragraph(lead_analysis)
    
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    filename = f"{company}_lead_analysis.docx"
    return filename, doc_stream

def create_full_docx(website_content: str, linkedin_content: str, news_content: str, company: str):
    doc = Document()
    doc.add_heading("Company Content and Analysis", 0)
    doc.add_heading("Company Website Content:", level=1)
    doc.add_paragraph(website_content)
    doc.add_heading("Company LinkedIn Profile Content:", level=1)
    doc.add_paragraph(linkedin_content)
    doc.add_heading("Recent News Articles:", level=1)
    doc.add_paragraph(news_content)
    
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    filename = f"{company}_full_content.docx"
    return filename, doc_stream

def send_email(access_token, recipient_emails, subject, body, attachments=None):
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    to_recipients = [{"emailAddress": {"address": mail}} for mail in recipient_emails]
    message = {
        "message": {
            "subject": subject,
            "body": {"contentType": "HTML", "content": body},
            "toRecipients": to_recipients,
            "attachments": []
        }
    }
    
    if attachments:
        for name, data_stream in attachments:
            encoded_content = base64.b64encode(data_stream.read()).decode("utf-8")
            message["message"]["attachments"].append({
                "@odata.type": "#microsoft.graph.fileAttachment",
                "name": name,
                "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "contentBytes": encoded_content
            })
    
    resp = requests.post(
        f"{GRAPH_API_ENDPOINT}/me/sendMail",
        headers=headers,
        json=message
    )
    if resp.status_code == 202:
        logging.info("✅ Email with attachments sent!")
        return True
    else:
        logging.error(f"❌ Failed to send: {resp.status_code} - {resp.text}")
        return False


def markdown_bold_to_html(text):
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)

def send_lead_data_to_api(
    lead_areas: str,
    account_name: str,
    lead_name: str,
    file_name: str | None = None,
    file_bytes: bytes | None = None,
) -> bool:
    """
    Sends lead data to the external API.
    - If file_name and file_bytes are provided, sends multipart/form-data using the exact pattern you confirmed works.
    - Otherwise, falls back to a JSON POST.
    """
    api_url = os.getenv("LEAD_API_URL")
    if not api_url:
        logging.error("Error: 'LEAD_API_URL' environment variable not found.")
        return False

    try:
        if file_name and file_bytes is not None:
            logging.info("Attempting to send lead data + file to API (multipart/form-data)...")

            files = [
                ('new_leadidentificationarea', (None, lead_areas)),
                ('new_name', (None, lead_name)),
                ('new_accountname', (None, account_name)),
                # Use the same field name you used in the working snippet:
                ('new_supportingdocuments', (file_name, io.BytesIO(file_bytes), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')),
            ]
            resp = requests.post(api_url, files=files, timeout=60)
        else:
            logging.info("Attempting to send lead data to API (JSON)...")
            data = {
                "new_leadidentificationarea": lead_areas,
                "new_name": lead_name,
                "new_accountname": account_name
            }
            resp = requests.post(api_url, json=data, timeout=60)

        resp.raise_for_status()
        logging.info("✅ Lead data successfully sent to API.")
        return True

    except requests.exceptions.RequestException as e:
        logging.error(f"❌ Failed to send lead data to API: {e}")
        if getattr(e, "response", None) is not None:
            logging.error(f"Status: {e.response.status_code}")
            logging.error(f"Response: {e.response.text}")
        return False

def extract_single_lead_details(lead_analysis, company):
    prompt = f"""
    Given the following lead analysis for the company "{company}", extract and return the following in plain text format:
    
    1. Customer Name
    2. Lead Identification Area (e.g., SFR150, Zones, DYN365, AI, AWS, Cost Takeout, Cloud Migration, Data Migration, Platform Migration, SaaS, GCC, Partner with IT) Mention any one. do not give more than one and no details.

    Format the output like this:
    
    **Customer Name**: <value>
    **Lead Identification Area**: <value>
    
    If a field is not found, just leave it like.

    Lead Analysis:
    {lead_analysis}
    """
    response = client.chat.completions.create(
        model=AZURE_DEPLOYMENT,
        messages=[{"role": "user", "content": prompt}],
    )
    extracted = response.choices[0].message.content.strip()
    return extracted

def check_potential_lead_by_area(single_lead_area, website_content, linkedin_content, news_content):
    """
    Analyzes content for a single, specific lead identification area.
    Returns the detailed analysis and a 'Yes' or 'No' based on that area.
    """
    combined_content = f"""
    You are a market-intelligence assistant. Your task is to analyze the following source materials and determine if there is any evidence of **{single_lead_area}**.

    Decision rules (strict):
    - Answer “Yes” ONLY if there is CLEAR, COMPANY-SPECIFIC evidence in the Sources below within the LAST 3 DAYS showing real interest or activity in **{single_lead_area}** (e.g., official announcements, press releases, case studies naming this company, exec statements, RFPs/tenders, partnerships, live projects, product/initiative pages tied to this company, or hiring/posts explicitly for this area).
    - Generic capability/marketing pages, industry articles not tied to this company, vague mentions, or items older than 3 days → “No”.
    - Use ONLY the provided Sources block. Do NOT invent or add outside links.
    - Every “Yes” MUST include at least one exact excerpt and a URL from the Sources. If you cannot provide an excerpt + URL, answer “No”.

    Provide a text report with the following sections:

    1. **{single_lead_area}**
       State "Yes" or "No". If "Yes", explain concisely why it is a potential lead and include an inline source link or supporting content if present. Do not add unrelated commentary.

    2. **Contacts**
       List all contacts found (email/phone/name+title) and the source source URL. Use only items present in Sources.
       - If no contacts are found, OMIT the **Contacts** section entirely. Do not write placeholders like "Not available".

    3. **Evidence**
       Under sub-headings **Website**, **LinkedIn**, and **News**, include ONLY those sub-headings that have at least one supporting excerpt. For each excerpt, include the source URL and, if present, the date.
       - If a source has no supporting excerpts, OMIT that sub-heading.
       - If there are zero excerpts across all sources, OMIT the entire **Evidence** section.
       - Never write placeholder lines like "No evidence found", "N/A", or similar.

    **Sources**
    Website Content:
    {website_content}

    LinkedIn Profile Content:
    {linkedin_content}

    Recent News Articles:
    {news_content}


    """
    
    response = client.chat.completions.create(
        model=AZURE_DEPLOYMENT,
        messages=[{"role": "user", "content": combined_content}],
    )
    analysis = response.choices[0].message.content
    
    # Check if the analysis confirms a lead for this specific area
    potential_lead_check = "Yes" if "Yes" in analysis or "yes" in analysis else "No"
    
    return analysis, potential_lead_check

# The Function App and Timer Trigger decorator
app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)
TARGET_COMPANY1 = os.getenv("TARGET_COMPANY1")


@app.function_name(name="ComputaCenter")
@app.schedule(schedule="0 30 10 3,6,9,12,15,18,21,24,27,30 * *", arg_name="myTimer", run_on_startup=False, use_monitor=True)
def ComputaCenter(myTimer: func.TimerRequest) -> None:
    utc_timestamp = datetime.utcnow()

    if myTimer.past_due:
        logging.warning("⏰ Timer is past due!")

    logging.info(f"🕒 Python timer trigger function started at: {utc_timestamp}")

    company = TARGET_COMPANY1
    pages = 1
    my_account_name = "Computacenter India"
    my_lead_name = "Lead from Lead Generator Tool"

    # 1. GNews Fetch
    logging.info(f"🔍 Fetching GNews articles for: {company}")
    today = datetime.utcnow()
    frm = (today - timedelta(days=30)).strftime("%Y-%m-%d")
    to = today.strftime("%Y-%m-%d")
    news_results = []
    try:
        resp = requests.get(
            BASE_URL,
            params={"q": company, "from": frm, "to": to, "lang": "en", "token": GNEWS_API_KEY},
            verify=False
        )
        resp.raise_for_status()
        for art in resp.json().get("articles", []):
            news_results.append({
                "title": art.get("title", ""),
                "description": art.get("description", ""),
                "url": art.get("url", "")
            })
        logging.info(f"✅ GNews API returned {len(news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ GNews API error: {e}")

    # 2. Google News Scrape
    logging.info("🌐 Scraping Google News...")
    try:
        google_news_results = scrape_google_news(company, pages)
        logging.info(f"✅ Google News scrape found {len(google_news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ Google News scraping error: {e}")
        google_news_results = []

    all_news = news_results + google_news_results
    if not all_news:
        logging.warning("⚠️ No news results found; skipping.")
        return

    news_content = "\n\n".join(
        f"**Title:** {n['title']}\n**Description:** {n['description']}\n**URL:** {n['url']}"
        for n in all_news
    )

    # 3. Website Scraping
    logging.info("🌐 Finding and scraping company website...")
    website = get_company_website(company, API_KEY, CX)
    if not website:
        logging.error(f"❌ Could not find website for {company}.")
        return
    try:
        website_content, _ = scrape_website(website)
        logging.info("✅ Website scraped successfully.")
    except Exception as e:
        logging.error(f"❌ Website scraping error: {e}")
        return

    # 4. LinkedIn Scraping
    logging.info("🔗 Finding and scraping LinkedIn profile...")
    linkedin_url = get_company_website(company + " LinkedIn", API_KEY, CX)
    if not linkedin_url:
        logging.error(f"❌ Could not find LinkedIn profile for {company}.")
        return
    try:
        linkedin_content, _ = scrape_website(linkedin_url)
        logging.info("✅ LinkedIn scraped successfully.")
    except Exception as e:
        logging.error(f"❌ LinkedIn scraping error: {e}")
        return
    # 5. AI Lead Check
    logging.info("🤖 Performing AI-based lead analysis...")
    try:
        lead_analysis, potential_lead_check = check_potential_lead(
            website_content, linkedin_content, news_content
        )
        logging.info("✅ Lead analysis complete.")
    except Exception as e:
        logging.error(f"❌ OpenAI analysis error: {e}")
        return
    logging.info(f"🔍 Is '{company}' a potential lead? → {potential_lead_check}")
    if potential_lead_check.strip().lower() == "yes": 
        logging.info("📌 Lead confirmed. Extracting all potential lead areas...") 
        
        # --- REVISED LOGIC FOR AREA EXTRACTION ---
        try:
            combined_details = extract_lead_details(lead_analysis, company)
            # This line will show you the exact string the AI returned.
            # Use this for debugging to see why the regex is not finding a match.
            logging.info(f"🔍 Raw AI output for lead areas: \n{combined_details}")
        except Exception as e:
            logging.error(f"❌ Failed to extract lead details from initial analysis: {e}")
            return
            
        all_lead_areas_list = []
        
        # Search for the "Lead Identification Area" line using a flexible regex
        m = re.search(r"\*\*Lead Identification Area\*\*: (.+)", combined_details)
        
        if m:
            all_lead_areas_str = m.group(1).strip()
            # Check if the AI returned a placeholder like "Not available"
            if all_lead_areas_str.lower() != "not available":
                all_lead_areas_list = [area.strip() for area in all_lead_areas_str.split(',') if area.strip()]

        if not all_lead_areas_list:
            logging.info("📝 The broad analysis did not identify any specific lead areas.")
            return

        logging.info(f"✅ Initially identified lead areas: {all_lead_areas_list}")

        # --- LOGIC TO FILTER OUT DUPLICATES BEFORE PROCESSING ---
        df = get_identified_leads_df()
        existing_company_row = df[df["Company Name"] == company]
        truly_new_areas_to_process = []

        if not existing_company_row.empty:
            existing_areas_str = existing_company_row["Lead Identification Areas"].iloc[0]
            existing_areas_set = set(normalize_areas_string(existing_areas_str).split(', ') if existing_areas_str else set())
            
            truly_new_areas_to_process = [area for area in all_lead_areas_list if area not in existing_areas_set]
            
            if not truly_new_areas_to_process:
                logging.info(f"📝 All identified areas for '{company}' are already in the Excel file. Skipping targeted analysis.")
                return
        else:
            truly_new_areas_to_process = all_lead_areas_list
        
        logging.info(f"✅ Found {len(truly_new_areas_to_process)} truly new areas to process: {truly_new_areas_to_process}")

        # --- LOOP OVER ONLY THE NEW AREAS ---

        for lead_area in truly_new_areas_to_process:
            logging.info(f"🔄 Re-analyzing content for specific new area: '{lead_area}'")
            try:
                area_analysis, area_potential_check = check_potential_lead_by_area(
                    lead_area, website_content, linkedin_content, news_content
                )
                
                if area_potential_check.strip().lower() == "yes":
                    area_details = extract_single_lead_details(area_analysis, company)
                    
                    email_flag = add_lead_to_excel(company, lead_area)
                    
                    if email_flag:
                        logging.info("✅ Excel updated successfully.")
                        
                        lead_doc_name, lead_doc_stream = create_lead_docx(area_analysis, company)
                        lead_doc_bytes = lead_doc_stream.getvalue() 
                        
                        logging.info("📄 DOCX files generated.")
                        
                        token = get_access_token() 
                        logging.info("🔐 Access token acquired.") 

                        email_body = ( 
                            f"<html><body><p>A new potential lead has been identified for <strong>{company}</strong> in the area of <strong>{lead_area}</strong>.</p>" 
                            + "".join(f"<p>{markdown_bold_to_html(line)}</p>" for line in area_details.splitlines()) 
                            + "<p>See attachments for full reports.</p></body></html>" 
                        ) 
                        
                        sent = send_email(
                            token,
                            ["vishnu.kg@sonata-software.com"],
                            f"New Lead: {company} - {lead_area}", 
                            email_body, 
                            attachments=[(lead_doc_name, lead_doc_stream)] 
                        ) 
                        
                        if sent: 
                            logging.info(f"✅ Email sent for area: '{lead_area}'") 
                            send_lead_data_to_api(lead_area, my_account_name, my_lead_name, lead_doc_name, lead_doc_bytes)
                            logging.info(f"📨 Lead data posted to external API for area: '{lead_area}'") 
                        else: 
                            logging.warning(f"⚠️ Email not sent for area: '{lead_area}'") 
                else:
                    logging.info(f"🚫 No lead indication found for area '{lead_area}'; skipping email and API steps.")
            except Exception as e:
                logging.error(f"❌ Error processing lead area '{lead_area}': {e}")
                
    else: 
        logging.info(f"🚫 '{company}' is not identified as a potential lead; skipping all downstream steps.") 

    logging.info("✅ Lead generation cycle completed.")


TARGET_COMPANY2 = os.getenv("TARGET_COMPANY2")
@app.function_name(name="PennyMac")
@app.schedule(schedule="0 35 10 3,6,9,12,15,18,21,24,27,30 * *", arg_name="myTimer", run_on_startup=False, use_monitor=True)
def PennyMac(myTimer: func.TimerRequest) -> None:
    utc_timestamp = datetime.utcnow()

    if myTimer.past_due:
        logging.warning("⏰ Timer is past due!")

    logging.info(f"🕒 Python timer trigger function started at: {utc_timestamp}")

    company = TARGET_COMPANY2
    pages = 1
    my_account_name = "PennyMac"
    my_lead_name = "Lead from Lead Generator Tool"

    # 1. GNews Fetch
    logging.info(f"🔍 Fetching GNews articles for: {company}")
    today = datetime.utcnow()
    frm = (today - timedelta(days=30)).strftime("%Y-%m-%d")
    to = today.strftime("%Y-%m-%d")
    news_results = []
    try:
        resp = requests.get(
            BASE_URL,
            params={"q": company, "from": frm, "to": to, "lang": "en", "token": GNEWS_API_KEY},
            verify=False
        )
        resp.raise_for_status()
        for art in resp.json().get("articles", []):
            news_results.append({
                "title": art.get("title", ""),
                "description": art.get("description", ""),
                "url": art.get("url", "")
            })
        logging.info(f"✅ GNews API returned {len(news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ GNews API error: {e}")

    # 2. Google News Scrape
    logging.info("🌐 Scraping Google News...")
    try:
        google_news_results = scrape_google_news(company, pages)
        logging.info(f"✅ Google News scrape found {len(google_news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ Google News scraping error: {e}")
        google_news_results = []

    all_news = news_results + google_news_results
    if not all_news:
        logging.warning("⚠️ No news results found; skipping.")
        return

    news_content = "\n\n".join(
        f"**Title:** {n['title']}\n**Description:** {n['description']}\n**URL:** {n['url']}"
        for n in all_news
    )

    # 3. Website Scraping
    logging.info("🌐 Finding and scraping company website...")
    website = get_company_website(company, API_KEY, CX)
    if not website:
        logging.error(f"❌ Could not find website for {company}.")
        return
    try:
        website_content, _ = scrape_website(website)
        logging.info("✅ Website scraped successfully.")
    except Exception as e:
        logging.error(f"❌ Website scraping error: {e}")
        return

    # 4. LinkedIn Scraping
    logging.info("🔗 Finding and scraping LinkedIn profile...")
    linkedin_url = get_company_website(company + " LinkedIn", API_KEY, CX)
    if not linkedin_url:
        logging.error(f"❌ Could not find LinkedIn profile for {company}.")
        return
    try:
        linkedin_content, _ = scrape_website(linkedin_url)
        logging.info("✅ LinkedIn scraped successfully.")
    except Exception as e:
        logging.error(f"❌ LinkedIn scraping error: {e}")
        return
    # 5. AI Lead Check
    logging.info("🤖 Performing AI-based lead analysis...")
    try:
        lead_analysis, potential_lead_check = check_potential_lead(
            website_content, linkedin_content, news_content
        )
        logging.info("✅ Lead analysis complete.")
    except Exception as e:
        logging.error(f"❌ OpenAI analysis error: {e}")
        return
    logging.info(f"🔍 Is '{company}' a potential lead? → {potential_lead_check}")
    if potential_lead_check.strip().lower() == "yes": 
        logging.info("📌 Lead confirmed. Extracting all potential lead areas...") 
        
        # --- REVISED LOGIC FOR AREA EXTRACTION ---
        try:
            combined_details = extract_lead_details(lead_analysis, company)
            # This line will show you the exact string the AI returned.
            # Use this for debugging to see why the regex is not finding a match.
            logging.info(f"🔍 Raw AI output for lead areas: \n{combined_details}")
        except Exception as e:
            logging.error(f"❌ Failed to extract lead details from initial analysis: {e}")
            return
            
        all_lead_areas_list = []
        
        # Search for the "Lead Identification Area" line using a flexible regex
        m = re.search(r"\*\*Lead Identification Area\*\*: (.+)", combined_details)
        
        if m:
            all_lead_areas_str = m.group(1).strip()
            # Check if the AI returned a placeholder like "Not available"
            if all_lead_areas_str.lower() != "not available":
                all_lead_areas_list = [area.strip() for area in all_lead_areas_str.split(',') if area.strip()]

        if not all_lead_areas_list:
            logging.info("📝 The broad analysis did not identify any specific lead areas.")
            return

        logging.info(f"✅ Initially identified lead areas: {all_lead_areas_list}")

        # --- LOGIC TO FILTER OUT DUPLICATES BEFORE PROCESSING ---
        df = get_identified_leads_df()
        existing_company_row = df[df["Company Name"] == company]
        truly_new_areas_to_process = []

        if not existing_company_row.empty:
            existing_areas_str = existing_company_row["Lead Identification Areas"].iloc[0]
            existing_areas_set = set(normalize_areas_string(existing_areas_str).split(', ') if existing_areas_str else set())
            
            truly_new_areas_to_process = [area for area in all_lead_areas_list if area not in existing_areas_set]
            
            if not truly_new_areas_to_process:
                logging.info(f"📝 All identified areas for '{company}' are already in the Excel file. Skipping targeted analysis.")
                return
        else:
            truly_new_areas_to_process = all_lead_areas_list
        
        logging.info(f"✅ Found {len(truly_new_areas_to_process)} truly new areas to process: {truly_new_areas_to_process}")

        # --- LOOP OVER ONLY THE NEW AREAS ---

        for lead_area in truly_new_areas_to_process:
            logging.info(f"🔄 Re-analyzing content for specific new area: '{lead_area}'")
            try:
                area_analysis, area_potential_check = check_potential_lead_by_area(
                    lead_area, website_content, linkedin_content, news_content
                )
                
                if area_potential_check.strip().lower() == "yes":
                    area_details = extract_single_lead_details(area_analysis, company)
                    
                    email_flag = add_lead_to_excel(company, lead_area)
                    
                    if email_flag:
                        logging.info("✅ Excel updated successfully.")
                        
                        lead_doc_name, lead_doc_stream = create_lead_docx(area_analysis, company)
                        lead_doc_bytes = lead_doc_stream.getvalue() 
                        
                        logging.info("📄 DOCX files generated.")
                        
                        token = get_access_token() 
                        logging.info("🔐 Access token acquired.") 

                        email_body = ( 
                            f"<html><body><p>A new potential lead has been identified for <strong>{company}</strong> in the area of <strong>{lead_area}</strong>.</p>" 
                            + "".join(f"<p>{markdown_bold_to_html(line)}</p>" for line in area_details.splitlines()) 
                            + "<p>See attachments for full reports.</p></body></html>" 
                        ) 
                        
                        sent = send_email(
                            token,
                            ["vishnu.kg@sonata-software.com"],
                            f"New Lead: {company} - {lead_area}", 
                            email_body, 
                            attachments=[(lead_doc_name, lead_doc_stream)] 
                        ) 
                        
                        if sent: 
                            logging.info(f"✅ Email sent for area: '{lead_area}'") 
                            send_lead_data_to_api(lead_area, my_account_name, my_lead_name, lead_doc_name, lead_doc_bytes)
                            logging.info(f"📨 Lead data posted to external API for area: '{lead_area}'") 
                        else: 
                            logging.warning(f"⚠️ Email not sent for area: '{lead_area}'") 
                else:
                    logging.info(f"🚫 No lead indication found for area '{lead_area}'; skipping email and API steps.")
            except Exception as e:
                logging.error(f"❌ Error processing lead area '{lead_area}': {e}")
                
    else: 
        logging.info(f"🚫 '{company}' is not identified as a potential lead; skipping all downstream steps.") 

    logging.info("✅ Lead generation cycle completed.")



TARGET_COMPANY3 = os.getenv("TARGET_COMPANY3")
@app.function_name(name="Fountaintire")
@app.schedule(schedule="0 40 10 3,6,9,12,15,18,21,24,27,30 * *", arg_name="myTimer", run_on_startup=False, use_monitor=True)
def Fountaintire(myTimer: func.TimerRequest) -> None:
    utc_timestamp = datetime.utcnow()

    if myTimer.past_due:
        logging.warning("⏰ Timer is past due!")

    logging.info(f"🕒 Python timer trigger function started at: {utc_timestamp}")

    company = TARGET_COMPANY3
    pages = 1
    my_account_name = "Fountain Tire"
    my_lead_name = "Lead from Lead Generator Tool"

    # 1. GNews Fetch
    logging.info(f"🔍 Fetching GNews articles for: {company}")
    today = datetime.utcnow()
    frm = (today - timedelta(days=30)).strftime("%Y-%m-%d")
    to = today.strftime("%Y-%m-%d")
    news_results = []
    try:
        resp = requests.get(
            BASE_URL,
            params={"q": company, "from": frm, "to": to, "lang": "en", "token": GNEWS_API_KEY},
            verify=False
        )
        resp.raise_for_status()
        for art in resp.json().get("articles", []):
            news_results.append({
                "title": art.get("title", ""),
                "description": art.get("description", ""),
                "url": art.get("url", "")
            })
        logging.info(f"✅ GNews API returned {len(news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ GNews API error: {e}")

    # 2. Google News Scrape
    logging.info("🌐 Scraping Google News...")
    try:
        google_news_results = scrape_google_news(company, pages)
        logging.info(f"✅ Google News scrape found {len(google_news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ Google News scraping error: {e}")
        google_news_results = []

    all_news = news_results + google_news_results
    if not all_news:
        logging.warning("⚠️ No news results found; skipping.")
        return

    news_content = "\n\n".join(
        f"**Title:** {n['title']}\n**Description:** {n['description']}\n**URL:** {n['url']}"
        for n in all_news
    )

    # 3. Website Scraping
    logging.info("🌐 Finding and scraping company website...")
    website = get_company_website(company, API_KEY, CX)
    if not website:
        logging.error(f"❌ Could not find website for {company}.")
        return
    try:
        website_content, _ = scrape_website(website)
        logging.info("✅ Website scraped successfully.")
    except Exception as e:
        logging.error(f"❌ Website scraping error: {e}")
        return

    # 4. LinkedIn Scraping
    logging.info("🔗 Finding and scraping LinkedIn profile...")
    linkedin_url = get_company_website(company + " LinkedIn", API_KEY, CX)
    if not linkedin_url:
        logging.error(f"❌ Could not find LinkedIn profile for {company}.")
        return
    try:
        linkedin_content, _ = scrape_website(linkedin_url)
        logging.info("✅ LinkedIn scraped successfully.")
    except Exception as e:
        logging.error(f"❌ LinkedIn scraping error: {e}")
        return
    # 5. AI Lead Check
    logging.info("🤖 Performing AI-based lead analysis...")
    try:
        lead_analysis, potential_lead_check = check_potential_lead(
            website_content, linkedin_content, news_content
        )
        logging.info("✅ Lead analysis complete.")
    except Exception as e:
        logging.error(f"❌ OpenAI analysis error: {e}")
        return
    logging.info(f"🔍 Is '{company}' a potential lead? → {potential_lead_check}")
    if potential_lead_check.strip().lower() == "yes": 
        logging.info("📌 Lead confirmed. Extracting all potential lead areas...") 
        
        # --- REVISED LOGIC FOR AREA EXTRACTION ---
        try:
            combined_details = extract_lead_details(lead_analysis, company)
            # This line will show you the exact string the AI returned.
            # Use this for debugging to see why the regex is not finding a match.
            logging.info(f"🔍 Raw AI output for lead areas: \n{combined_details}")
        except Exception as e:
            logging.error(f"❌ Failed to extract lead details from initial analysis: {e}")
            return
            
        all_lead_areas_list = []
        
        # Search for the "Lead Identification Area" line using a flexible regex
        m = re.search(r"\*\*Lead Identification Area\*\*: (.+)", combined_details)
        
        if m:
            all_lead_areas_str = m.group(1).strip()
            # Check if the AI returned a placeholder like "Not available"
            if all_lead_areas_str.lower() != "not available":
                all_lead_areas_list = [area.strip() for area in all_lead_areas_str.split(',') if area.strip()]

        if not all_lead_areas_list:
            logging.info("📝 The broad analysis did not identify any specific lead areas.")
            return

        logging.info(f"✅ Initially identified lead areas: {all_lead_areas_list}")

        # --- LOGIC TO FILTER OUT DUPLICATES BEFORE PROCESSING ---
        df = get_identified_leads_df()
        existing_company_row = df[df["Company Name"] == company]
        truly_new_areas_to_process = []

        if not existing_company_row.empty:
            existing_areas_str = existing_company_row["Lead Identification Areas"].iloc[0]
            existing_areas_set = set(normalize_areas_string(existing_areas_str).split(', ') if existing_areas_str else set())
            
            truly_new_areas_to_process = [area for area in all_lead_areas_list if area not in existing_areas_set]
            
            if not truly_new_areas_to_process:
                logging.info(f"📝 All identified areas for '{company}' are already in the Excel file. Skipping targeted analysis.")
                return
        else:
            truly_new_areas_to_process = all_lead_areas_list
        
        logging.info(f"✅ Found {len(truly_new_areas_to_process)} truly new areas to process: {truly_new_areas_to_process}")

        # --- LOOP OVER ONLY THE NEW AREAS ---

        for lead_area in truly_new_areas_to_process:
            logging.info(f"🔄 Re-analyzing content for specific new area: '{lead_area}'")
            try:
                area_analysis, area_potential_check = check_potential_lead_by_area(
                    lead_area, website_content, linkedin_content, news_content
                )
                
                if area_potential_check.strip().lower() == "yes":
                    area_details = extract_single_lead_details(area_analysis, company)
                    
                    email_flag = add_lead_to_excel(company, lead_area)
                    
                    if email_flag:
                        logging.info("✅ Excel updated successfully.")
                        
                        lead_doc_name, lead_doc_stream = create_lead_docx(area_analysis, company)
                        lead_doc_bytes = lead_doc_stream.getvalue() 
                        
                        logging.info("📄 DOCX files generated.")
                        
                        token = get_access_token() 
                        logging.info("🔐 Access token acquired.") 

                        email_body = ( 
                            f"<html><body><p>A new potential lead has been identified for <strong>{company}</strong> in the area of <strong>{lead_area}</strong>.</p>" 
                            + "".join(f"<p>{markdown_bold_to_html(line)}</p>" for line in area_details.splitlines()) 
                            + "<p>See attachments for full reports.</p></body></html>" 
                        ) 
                        
                        sent = send_email(
                            token,
                            ["vishnu.kg@sonata-software.com"],
                            f"New Lead: {company} - {lead_area}", 
                            email_body, 
                            attachments=[(lead_doc_name, lead_doc_stream)] 
                        ) 
                        
                        if sent: 
                            logging.info(f"✅ Email sent for area: '{lead_area}'") 
                            send_lead_data_to_api(lead_area, my_account_name, my_lead_name, lead_doc_name, lead_doc_bytes)
                            logging.info(f"📨 Lead data posted to external API for area: '{lead_area}'") 
                        else: 
                            logging.warning(f"⚠️ Email not sent for area: '{lead_area}'") 
                else:
                    logging.info(f"🚫 No lead indication found for area '{lead_area}'; skipping email and API steps.")
            except Exception as e:
                logging.error(f"❌ Error processing lead area '{lead_area}': {e}")
                
    else: 
        logging.info(f"🚫 '{company}' is not identified as a potential lead; skipping all downstream steps.") 

    logging.info("✅ Lead generation cycle completed.")


TARGET_COMPANY4 = os.getenv("TARGET_COMPANY4")
@app.function_name(name="Wellpath")
@app.schedule(schedule="0 45 10 3,6,9,12,15,18,21,24,27,30 * *", arg_name="myTimer", run_on_startup=False, use_monitor=True)
def Wellpath(myTimer: func.TimerRequest) -> None:
    utc_timestamp = datetime.utcnow()

    if myTimer.past_due:
        logging.warning("⏰ Timer is past due!")

    logging.info(f"🕒 Python timer trigger function started at: {utc_timestamp}")

    company = TARGET_COMPANY4
    pages = 1
    my_account_name = "Wellpath"
    my_lead_name = "Lead from Lead Generator Tool"

    # 1. GNews Fetch
    logging.info(f"🔍 Fetching GNews articles for: {company}")
    today = datetime.utcnow()
    frm = (today - timedelta(days=30)).strftime("%Y-%m-%d")
    to = today.strftime("%Y-%m-%d")
    news_results = []
    try:
        resp = requests.get(
            BASE_URL,
            params={"q": company, "from": frm, "to": to, "lang": "en", "token": GNEWS_API_KEY},
            verify=False
        )
        resp.raise_for_status()
        for art in resp.json().get("articles", []):
            news_results.append({
                "title": art.get("title", ""),
                "description": art.get("description", ""),
                "url": art.get("url", "")
            })
        logging.info(f"✅ GNews API returned {len(news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ GNews API error: {e}")

    # 2. Google News Scrape
    logging.info("🌐 Scraping Google News...")
    try:
        google_news_results = scrape_google_news(company, pages)
        logging.info(f"✅ Google News scrape found {len(google_news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ Google News scraping error: {e}")
        google_news_results = []

    all_news = news_results + google_news_results
    if not all_news:
        logging.warning("⚠️ No news results found; skipping.")
        return

    news_content = "\n\n".join(
        f"**Title:** {n['title']}\n**Description:** {n['description']}\n**URL:** {n['url']}"
        for n in all_news
    )

    # 3. Website Scraping
    logging.info("🌐 Finding and scraping company website...")
    website = get_company_website(company, API_KEY, CX)
    if not website:
        logging.error(f"❌ Could not find website for {company}.")
        return
    try:
        website_content, _ = scrape_website(website)
        logging.info("✅ Website scraped successfully.")
    except Exception as e:
        logging.error(f"❌ Website scraping error: {e}")
        return

    # 4. LinkedIn Scraping
    logging.info("🔗 Finding and scraping LinkedIn profile...")
    linkedin_url = get_company_website(company + " LinkedIn", API_KEY, CX)
    if not linkedin_url:
        logging.error(f"❌ Could not find LinkedIn profile for {company}.")
        return
    try:
        linkedin_content, _ = scrape_website(linkedin_url)
        logging.info("✅ LinkedIn scraped successfully.")
    except Exception as e:
        logging.error(f"❌ LinkedIn scraping error: {e}")
        return
    # 5. AI Lead Check
    logging.info("🤖 Performing AI-based lead analysis...")
    try:
        lead_analysis, potential_lead_check = check_potential_lead(
            website_content, linkedin_content, news_content
        )
        logging.info("✅ Lead analysis complete.")
    except Exception as e:
        logging.error(f"❌ OpenAI analysis error: {e}")
        return
    logging.info(f"🔍 Is '{company}' a potential lead? → {potential_lead_check}")
    if potential_lead_check.strip().lower() == "yes": 
        logging.info("📌 Lead confirmed. Extracting all potential lead areas...") 
        
        # --- REVISED LOGIC FOR AREA EXTRACTION ---
        try:
            combined_details = extract_lead_details(lead_analysis, company)
            # This line will show you the exact string the AI returned.
            # Use this for debugging to see why the regex is not finding a match.
            logging.info(f"🔍 Raw AI output for lead areas: \n{combined_details}")
        except Exception as e:
            logging.error(f"❌ Failed to extract lead details from initial analysis: {e}")
            return
            
        all_lead_areas_list = []
        
        # Search for the "Lead Identification Area" line using a flexible regex
        m = re.search(r"\*\*Lead Identification Area\*\*: (.+)", combined_details)
        
        if m:
            all_lead_areas_str = m.group(1).strip()
            # Check if the AI returned a placeholder like "Not available"
            if all_lead_areas_str.lower() != "not available":
                all_lead_areas_list = [area.strip() for area in all_lead_areas_str.split(',') if area.strip()]

        if not all_lead_areas_list:
            logging.info("📝 The broad analysis did not identify any specific lead areas.")
            return

        logging.info(f"✅ Initially identified lead areas: {all_lead_areas_list}")

        # --- LOGIC TO FILTER OUT DUPLICATES BEFORE PROCESSING ---
        df = get_identified_leads_df()
        existing_company_row = df[df["Company Name"] == company]
        truly_new_areas_to_process = []

        if not existing_company_row.empty:
            existing_areas_str = existing_company_row["Lead Identification Areas"].iloc[0]
            existing_areas_set = set(normalize_areas_string(existing_areas_str).split(', ') if existing_areas_str else set())
            
            truly_new_areas_to_process = [area for area in all_lead_areas_list if area not in existing_areas_set]
            
            if not truly_new_areas_to_process:
                logging.info(f"📝 All identified areas for '{company}' are already in the Excel file. Skipping targeted analysis.")
                return
        else:
            truly_new_areas_to_process = all_lead_areas_list
        
        logging.info(f"✅ Found {len(truly_new_areas_to_process)} truly new areas to process: {truly_new_areas_to_process}")

        # --- LOOP OVER ONLY THE NEW AREAS ---

        for lead_area in truly_new_areas_to_process:
            logging.info(f"🔄 Re-analyzing content for specific new area: '{lead_area}'")
            try:
                area_analysis, area_potential_check = check_potential_lead_by_area(
                    lead_area, website_content, linkedin_content, news_content
                )
                
                if area_potential_check.strip().lower() == "yes":
                    area_details = extract_single_lead_details(area_analysis, company)
                    
                    email_flag = add_lead_to_excel(company, lead_area)
                    
                    if email_flag:
                        logging.info("✅ Excel updated successfully.")
                        
                        lead_doc_name, lead_doc_stream = create_lead_docx(area_analysis, company)
                        lead_doc_bytes = lead_doc_stream.getvalue() 
                        
                        logging.info("📄 DOCX files generated.")
                        
                        token = get_access_token() 
                        logging.info("🔐 Access token acquired.") 

                        email_body = ( 
                            f"<html><body><p>A new potential lead has been identified for <strong>{company}</strong> in the area of <strong>{lead_area}</strong>.</p>" 
                            + "".join(f"<p>{markdown_bold_to_html(line)}</p>" for line in area_details.splitlines()) 
                            + "<p>See attachments for full reports.</p></body></html>" 
                        ) 
                        
                        sent = send_email(
                            token,
                            ["vishnu.kg@sonata-software.com"],
                            f"New Lead: {company} - {lead_area}", 
                            email_body, 
                            attachments=[(lead_doc_name, lead_doc_stream)] 
                        ) 
                        
                        if sent: 
                            logging.info(f"✅ Email sent for area: '{lead_area}'") 
                            send_lead_data_to_api(lead_area, my_account_name, my_lead_name, lead_doc_name, lead_doc_bytes)
                            logging.info(f"📨 Lead data posted to external API for area: '{lead_area}'") 
                        else: 
                            logging.warning(f"⚠️ Email not sent for area: '{lead_area}'") 
                else:
                    logging.info(f"🚫 No lead indication found for area '{lead_area}'; skipping email and API steps.")
            except Exception as e:
                logging.error(f"❌ Error processing lead area '{lead_area}': {e}")
                
    else: 
        logging.info(f"🚫 '{company}' is not identified as a potential lead; skipping all downstream steps.") 

    logging.info("✅ Lead generation cycle completed.")

TARGET_COMPANY5 = os.getenv("TARGET_COMPANY5")
@app.function_name(name="TUI")
@app.schedule(schedule="0 50 10 3,6,9,12,15,18,21,24,27,30 * *", arg_name="myTimer", run_on_startup=False, use_monitor=True)
def TUI(myTimer: func.TimerRequest) -> None:
    utc_timestamp = datetime.utcnow()

    if myTimer.past_due:
        logging.warning("⏰ Timer is past due!")

    logging.info(f"🕒 Python timer trigger function started at: {utc_timestamp}")

    company = TARGET_COMPANY5
    pages = 1
    my_account_name = "TUI"
    my_lead_name = "Lead from Lead Generator Tool"

    # 1. GNews Fetch
    logging.info(f"🔍 Fetching GNews articles for: {company}")
    today = datetime.utcnow()
    frm = (today - timedelta(days=30)).strftime("%Y-%m-%d")
    to = today.strftime("%Y-%m-%d")
    news_results = []
    try:
        resp = requests.get(
            BASE_URL,
            params={"q": company, "from": frm, "to": to, "lang": "en", "token": GNEWS_API_KEY},
            verify=False
        )
        resp.raise_for_status()
        for art in resp.json().get("articles", []):
            news_results.append({
                "title": art.get("title", ""),
                "description": art.get("description", ""),
                "url": art.get("url", "")
            })
        logging.info(f"✅ GNews API returned {len(news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ GNews API error: {e}")

    # 2. Google News Scrape
    logging.info("🌐 Scraping Google News...")
    try:
        google_news_results = scrape_google_news(company, pages)
        logging.info(f"✅ Google News scrape found {len(google_news_results)} articles.")
    except Exception as e:
        logging.error(f"❌ Google News scraping error: {e}")
        google_news_results = []

    all_news = news_results + google_news_results
    if not all_news:
        logging.warning("⚠️ No news results found; skipping.")
        return

    news_content = "\n\n".join(
        f"**Title:** {n['title']}\n**Description:** {n['description']}\n**URL:** {n['url']}"
        for n in all_news
    )

    # 3. Website Scraping
    logging.info("🌐 Finding and scraping company website...")
    website = get_company_website(company, API_KEY, CX)
    if not website:
        logging.error(f"❌ Could not find website for {company}.")
        return
    try:
        website_content, _ = scrape_website(website)
        logging.info("✅ Website scraped successfully.")
    except Exception as e:
        logging.error(f"❌ Website scraping error: {e}")
        return

    # 4. LinkedIn Scraping
    logging.info("🔗 Finding and scraping LinkedIn profile...")
    linkedin_url = get_company_website(company + " LinkedIn", API_KEY, CX)
    if not linkedin_url:
        logging.error(f"❌ Could not find LinkedIn profile for {company}.")
        return
    try:
        linkedin_content, _ = scrape_website(linkedin_url)
        logging.info("✅ LinkedIn scraped successfully.")
    except Exception as e:
        logging.error(f"❌ LinkedIn scraping error: {e}")
        return
    # 5. AI Lead Check
    logging.info("🤖 Performing AI-based lead analysis...")
    try:
        lead_analysis, potential_lead_check = check_potential_lead(
            website_content, linkedin_content, news_content
        )
        logging.info("✅ Lead analysis complete.")
    except Exception as e:
        logging.error(f"❌ OpenAI analysis error: {e}")
        return
    logging.info(f"🔍 Is '{company}' a potential lead? → {potential_lead_check}")
    if potential_lead_check.strip().lower() == "yes": 
        logging.info("📌 Lead confirmed. Extracting all potential lead areas...") 
        
        # --- REVISED LOGIC FOR AREA EXTRACTION ---
        try:
            combined_details = extract_lead_details(lead_analysis, company)
            # This line will show you the exact string the AI returned.
            # Use this for debugging to see why the regex is not finding a match.
            logging.info(f"🔍 Raw AI output for lead areas: \n{combined_details}")
        except Exception as e:
            logging.error(f"❌ Failed to extract lead details from initial analysis: {e}")
            return
            
        all_lead_areas_list = []
        
        # Search for the "Lead Identification Area" line using a flexible regex
        m = re.search(r"\*\*Lead Identification Area\*\*: (.+)", combined_details)
        
        if m:
            all_lead_areas_str = m.group(1).strip()
            # Check if the AI returned a placeholder like "Not available"
            if all_lead_areas_str.lower() != "not available":
                all_lead_areas_list = [area.strip() for area in all_lead_areas_str.split(',') if area.strip()]

        if not all_lead_areas_list:
            logging.info("📝 The broad analysis did not identify any specific lead areas.")
            return

        logging.info(f"✅ Initially identified lead areas: {all_lead_areas_list}")

        # --- LOGIC TO FILTER OUT DUPLICATES BEFORE PROCESSING ---
        df = get_identified_leads_df()
        existing_company_row = df[df["Company Name"] == company]
        truly_new_areas_to_process = []

        if not existing_company_row.empty:
            existing_areas_str = existing_company_row["Lead Identification Areas"].iloc[0]
            existing_areas_set = set(normalize_areas_string(existing_areas_str).split(', ') if existing_areas_str else set())
            
            truly_new_areas_to_process = [area for area in all_lead_areas_list if area not in existing_areas_set]
            
            if not truly_new_areas_to_process:
                logging.info(f"📝 All identified areas for '{company}' are already in the Excel file. Skipping targeted analysis.")
                return
        else:
            truly_new_areas_to_process = all_lead_areas_list
        
        logging.info(f"✅ Found {len(truly_new_areas_to_process)} truly new areas to process: {truly_new_areas_to_process}")

        # --- LOOP OVER ONLY THE NEW AREAS ---

        for lead_area in truly_new_areas_to_process:
            logging.info(f"🔄 Re-analyzing content for specific new area: '{lead_area}'")
            try:
                area_analysis, area_potential_check = check_potential_lead_by_area(
                    lead_area, website_content, linkedin_content, news_content
                )
                
                if area_potential_check.strip().lower() == "yes":
                    area_details = extract_single_lead_details(area_analysis, company)
                    
                    email_flag = add_lead_to_excel(company, lead_area)
                    
                    if email_flag:
                        logging.info("✅ Excel updated successfully.")
                        
                        lead_doc_name, lead_doc_stream = create_lead_docx(area_analysis, company)
                        lead_doc_bytes = lead_doc_stream.getvalue() 
                        
                        logging.info("📄 DOCX files generated.")
                        
                        token = get_access_token() 
                        logging.info("🔐 Access token acquired.") 

                        email_body = ( 
                            f"<html><body><p>A new potential lead has been identified for <strong>{company}</strong> in the area of <strong>{lead_area}</strong>.</p>" 
                            + "".join(f"<p>{markdown_bold_to_html(line)}</p>" for line in area_details.splitlines()) 
                            + "<p>See attachments for full reports.</p></body></html>" 
                        ) 
                        
                        sent = send_email(
                            token,
                            ["vishnu.kg@sonata-software.com"],
                            f"New Lead: {company} - {lead_area}", 
                            email_body, 
                            attachments=[(lead_doc_name, lead_doc_stream)] 
                        ) 
                        
                        if sent: 
                            logging.info(f"✅ Email sent for area: '{lead_area}'") 
                            send_lead_data_to_api(lead_area, my_account_name, my_lead_name, lead_doc_name, lead_doc_bytes)
                            logging.info(f"📨 Lead data posted to external API for area: '{lead_area}'") 
                        else: 
                            logging.warning(f"⚠️ Email not sent for area: '{lead_area}'") 
                else:
                    logging.info(f"🚫 No lead indication found for area '{lead_area}'; skipping email and API steps.")
            except Exception as e:
                logging.error(f"❌ Error processing lead area '{lead_area}': {e}")
                
    else: 
        logging.info(f"🚫 '{company}' is not identified as a potential lead; skipping all downstream steps.") 

    logging.info("✅ Lead generation cycle completed.")
